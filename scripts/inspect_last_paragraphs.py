import docx

doc_path = r"d:\Workplace\SummVi\docs\Báo cáo (1).docx"
doc = docx.Document(doc_path)

with open("inspect_end_results.txt", "w", encoding="utf-8") as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
    for idx in range(740, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        f.write(f"Index: {idx} | Style: {para.style.name} | Runs: {len(para.runs)} | Text: '{para.text}'\n")

print("Done!")
