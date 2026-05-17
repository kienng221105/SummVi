import docx

doc_path = r"d:\Workplace\SummVi\docs\Báo cáo (1).docx"
doc = docx.Document(doc_path)

with open("inspect_results.txt", "w", encoding="utf-8") as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(keyword in text for keyword in ["Kết luận tổng hợp", "Hướng phát triển"]):
            f.write(f"Index: {idx} | Style: {para.style.name} | Text: '{text}'\n")

print("Done! Results written to inspect_results.txt")
