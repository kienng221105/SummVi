import docx
import os

def extract_docx_to_md(docx_path, md_path):
    print(f"Reading from {docx_path}...")
    doc = docx.Document(docx_path)
    print(f"Loaded document. Paragraphs count: {len(doc.paragraphs)}. Tables count: {len(doc.tables)}")
    
    with open(md_path, "w", encoding="utf-8") as f:
        # We will iterate through child elements of document body in order if possible, 
        # but iterating through paragraphs is a simple and reliable start.
        for para in doc.paragraphs:
            style = para.style.name
            text = para.text.strip()
            if not text:
                continue
            
            if style.startswith('Heading 1'):
                f.write(f"\n# {text}\n\n")
            elif style.startswith('Heading 2'):
                f.write(f"\n## {text}\n\n")
            elif style.startswith('Heading 3'):
                f.write(f"\n### {text}\n\n")
            elif style.startswith('Heading 4'):
                f.write(f"\n#### {text}\n\n")
            else:
                if 'List Bullet' in style:
                    f.write(f"- {text}\n")
                elif 'List' in style:
                    f.write(f"1. {text}\n")
                else:
                    f.write(f"{text}\n\n")
                    
        # Let's also extract table text and append at the end or print some summary of tables
        if doc.tables:
            f.write("\n\n---\n# TABLES SUMMARY\n\n")
            for idx, table in enumerate(doc.tables):
                f.write(f"### Table {idx + 1}\n\n")
                for row in table.rows:
                    cells_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    f.write("| " + " | ".join(cells_text) + " |\n")
                f.write("\n")

    print(f"Markdown written successfully to {md_path}!")

if __name__ == "__main__":
    docx_path = r"d:\Workplace\SummVi\docs\Báo cáo.docx"
    md_path = r"d:\Workplace\SummVi\docs\Bao_cao_hoan_thanh.md"
    extract_docx_to_md(docx_path, md_path)
