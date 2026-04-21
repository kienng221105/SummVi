import fitz
from docx import Document as DocxDocument
from .schema import PageData
def extract_pdf(file_path: str) -> list[PageData]:
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        images = page.get_images(full=True)
        pages.append(PageData(
            page_num=i + 1,
            text=text,
            char_count=len(text.strip()),
            has_images=len(images) > 0,
        ))
    doc.close()
    return pages

def extract_pdf_images(file_path: str, page_num: int) -> list[bytes]:
    doc = fitz.open(file_path)
    page = doc[page_num - 1]
    image_bytes = []
    for img in page.get_images(full=True):
        xref = img[0]
        base_image = doc.extract_image(xref)
        image_bytes.append(base_image["image"])
    doc.close()
    return image_bytes

def extract_pdf_page_as_image(file_path: str, page_num: int, dpi: int = 300) -> bytes:
    doc = fitz.open(file_path)
    page = doc[page_num - 1]
    pix = page.get_pixmap(dpi=dpi)
    img_bytes = pix.tobytes("png")
    doc.close()
    return img_bytes

def extract_pdf_metadata(file_path: str) -> dict:
    doc = fitz.open(file_path)
    meta = doc.metadata
    doc.close()
    return meta or {}

def extract_docx(file_path: str) -> list[PageData]:
    doc = DocxDocument(file_path)
    chunks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        prefix = _heading_prefix(style)
        chunks.append(f"{prefix}{text}")
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                chunks.append(f"[TABLE] {row_text}")
    full_text = "\n".join(chunks)
    return [PageData(
        page_num=1,
        text=full_text,
        char_count=len(full_text.strip()),
        has_images=False,
    )]

def extract_docx_metadata(file_path: str) -> dict:
    doc = DocxDocument(file_path)
    props = doc.core_properties
    return {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created or ""),
    }

def _heading_prefix(style_name: str) -> str:
    mapping = {
        "Heading 1": "# ",
        "Heading 2": "## ",
        "Heading 3": "### ",
    }
    return mapping.get(style_name, "")